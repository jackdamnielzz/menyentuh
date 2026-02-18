#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import shutil
import sys
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


TOP_LEVEL_REQUIRED_KEYS = ("metadata", "branding", "document_layout", "sections")
FIELD_REQUIRED_KEYS = ("id", "label", "type", "required", "help_text")
ALLOWED_FIELD_TYPES = {
    "text_short",
    "text_long",
    "checkbox",
    "checkbox_group",
    "dropdown",
    "date",
    "signature_line",
    "info_block",
}

WD_CONTENT_CONTROL_RICH_TEXT = 0
WD_CONTENT_CONTROL_TEXT = 1
WD_CONTENT_CONTROL_DROPDOWN = 4
WD_CONTENT_CONTROL_DATE = 6
WD_CONTENT_CONTROL_CHECKBOX = 8
WD_FIND_STOP = 0


@dataclass(frozen=True)
class PlaceholderSpec:
    marker: str
    field_type: str
    field_id: str
    label: str
    required: bool
    options: list[str] = field(default_factory=list)


class SpecValidationError(RuntimeError):
    pass


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Genereer een invulbaar Menyentuh intake/anamneseformulier in Word (.docx)."
    )
    parser.add_argument(
        "--spec",
        default="forms/intake_anamnese_spec.json",
        help="Pad naar de JSON-formulierspecificatie.",
    )
    parser.add_argument(
        "--out",
        default="output/Menyentuh_Intake_Anamnese.docx",
        help="Pad van het output Word-bestand.",
    )
    parser.add_argument(
        "--logo",
        default=None,
        help="Optioneel pad naar een logo, overschrijft branding.logo_path uit de spec.",
    )
    parser.add_argument(
        "--validate-only",
        action="store_true",
        help="Voer alleen validatie uit op de spec en stop daarna.",
    )
    parser.add_argument(
        "--no-com",
        action="store_true",
        help="Sla COM-verrijking over en maak alleen een niet-interactieve fallbackversie.",
    )
    return parser.parse_args()


def load_spec(spec_path: Path) -> dict[str, Any]:
    if not spec_path.exists():
        raise FileNotFoundError(f"Spec-bestand niet gevonden: {spec_path}")
    try:
        return json.loads(spec_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise SpecValidationError(f"Ongeldige JSON in {spec_path}: {exc}") from exc


def validate_spec(spec: dict[str, Any]) -> None:
    errors: list[str] = []

    if not isinstance(spec, dict):
        raise SpecValidationError("Spec-root moet een JSON-object zijn.")

    for key in TOP_LEVEL_REQUIRED_KEYS:
        if key not in spec:
            errors.append(f"Ontbrekende top-level key: '{key}'.")

    if errors:
        raise SpecValidationError("\n".join(errors))

    metadata = spec["metadata"]
    branding = spec["branding"]
    layout = spec["document_layout"]
    sections = spec["sections"]

    if not isinstance(metadata, dict):
        errors.append("'metadata' moet een object zijn.")
    if not isinstance(branding, dict):
        errors.append("'branding' moet een object zijn.")
    if not isinstance(layout, dict):
        errors.append("'document_layout' moet een object zijn.")
    if not isinstance(sections, list) or not sections:
        errors.append("'sections' moet een niet-lege lijst zijn.")

    if isinstance(layout, dict):
        margin_cm = layout.get("margin_cm", 2.0)
        if not isinstance(margin_cm, (int, float)) or margin_cm <= 0:
            errors.append("'document_layout.margin_cm' moet een positief getal zijn.")

    if isinstance(branding, dict):
        for color_key in (
            "primary_color",
            "secondary_color",
            "accent_color",
            "surface_color",
            "text_color",
        ):
            if color_key not in branding or not is_hex_color(str(branding.get(color_key, ""))):
                errors.append(f"'branding.{color_key}' moet een hex kleur zijn, bv. #17372c.")
        if not isinstance(branding.get("heading_font", ""), str) or not branding["heading_font"]:
            errors.append("'branding.heading_font' moet een niet-lege string zijn.")
        if not isinstance(branding.get("body_font", ""), str) or not branding["body_font"]:
            errors.append("'branding.body_font' moet een niet-lege string zijn.")
        if not isinstance(branding.get("logo_path", ""), str) or not branding["logo_path"]:
            errors.append("'branding.logo_path' moet een niet-lege string zijn.")

    seen_ids: set[str] = set()

    if isinstance(sections, list):
        for section_index, section in enumerate(sections, start=1):
            section_prefix = f"sections[{section_index}]"
            if not isinstance(section, dict):
                errors.append(f"{section_prefix} moet een object zijn.")
                continue

            for section_key in ("id", "title", "fields"):
                if section_key not in section:
                    errors.append(f"{section_prefix}: ontbrekende key '{section_key}'.")

            if not isinstance(section.get("id"), str) or not section["id"].strip():
                errors.append(f"{section_prefix}.id moet een niet-lege string zijn.")
            if not isinstance(section.get("title"), str) or not section["title"].strip():
                errors.append(f"{section_prefix}.title moet een niet-lege string zijn.")

            fields = section.get("fields")
            if not isinstance(fields, list) or not fields:
                errors.append(f"{section_prefix}.fields moet een niet-lege lijst zijn.")
                continue

            for field_index, field_spec in enumerate(fields, start=1):
                field_prefix = f"{section_prefix}.fields[{field_index}]"
                if not isinstance(field_spec, dict):
                    errors.append(f"{field_prefix} moet een object zijn.")
                    continue

                for field_key in FIELD_REQUIRED_KEYS:
                    if field_key not in field_spec:
                        errors.append(f"{field_prefix}: ontbrekende key '{field_key}'.")

                if any(field_key not in field_spec for field_key in FIELD_REQUIRED_KEYS):
                    continue

                field_id = field_spec["id"]
                field_type = field_spec["type"]
                required = field_spec["required"]
                help_text = field_spec["help_text"]

                if not isinstance(field_id, str) or not field_id.strip():
                    errors.append(f"{field_prefix}.id moet een niet-lege string zijn.")
                    continue

                if field_id in seen_ids:
                    errors.append(f"Dubbele field id gevonden: '{field_id}'.")
                else:
                    seen_ids.add(field_id)

                if field_type not in ALLOWED_FIELD_TYPES:
                    errors.append(
                        f"{field_prefix}.type '{field_type}' is ongeldig. Toegestaan: {sorted(ALLOWED_FIELD_TYPES)}."
                    )

                if not isinstance(required, bool):
                    errors.append(f"{field_prefix}.required moet true of false zijn.")

                if not isinstance(field_spec["label"], str) or not field_spec["label"].strip():
                    errors.append(f"{field_prefix}.label moet een niet-lege string zijn.")

                if not isinstance(help_text, str):
                    errors.append(f"{field_prefix}.help_text moet een string zijn (leeg is toegestaan).")

                if field_type == "dropdown":
                    options = field_spec.get("options")
                    if not isinstance(options, list) or not options:
                        errors.append(f"{field_prefix}.options moet een niet-lege lijst zijn voor dropdown.")
                    elif not all(isinstance(option, str) and option.strip() for option in options):
                        errors.append(f"{field_prefix}.options mag alleen niet-lege strings bevatten.")

                if field_type == "checkbox_group":
                    options = field_spec.get("options")
                    if not isinstance(options, list) or not options:
                        errors.append(f"{field_prefix}.options moet een niet-lege lijst zijn voor checkbox_group.")
                    else:
                        for option_index, option in enumerate(options, start=1):
                            option_prefix = f"{field_prefix}.options[{option_index}]"
                            if not isinstance(option, dict):
                                errors.append(f"{option_prefix} moet een object zijn.")
                                continue
                            if "id" not in option or "label" not in option:
                                errors.append(f"{option_prefix} mist 'id' en/of 'label'.")
                                continue
                            option_id = option["id"]
                            option_label = option["label"]
                            if not isinstance(option_id, str) or not option_id.strip():
                                errors.append(f"{option_prefix}.id moet een niet-lege string zijn.")
                                continue
                            if option_id in seen_ids:
                                errors.append(f"Dubbele option id gevonden: '{option_id}'.")
                            else:
                                seen_ids.add(option_id)
                            if not isinstance(option_label, str) or not option_label.strip():
                                errors.append(f"{option_prefix}.label moet een niet-lege string zijn.")

    if errors:
        raise SpecValidationError("\n".join(errors))


def is_hex_color(value: str) -> bool:
    stripped = value.strip()
    if stripped.startswith("#"):
        stripped = stripped[1:]
    if len(stripped) != 6:
        return False
    return all(char in "0123456789abcdefABCDEF" for char in stripped)


def clean_hex(value: str, fallback: str) -> str:
    source = value if is_hex_color(value) else fallback
    source = source.strip()
    if source.startswith("#"):
        source = source[1:]
    return source.upper()


def hex_to_rgb(value: str, fallback: str = "#000000") -> RGBColor:
    hex_value = clean_hex(value, fallback)
    return RGBColor(int(hex_value[0:2], 16), int(hex_value[2:4], 16), int(hex_value[4:6], 16))


def marker_for(field_id: str) -> str:
    return f"[[FIELD:{field_id}]]"


def resolve_logo_path(spec_path: Path, spec: dict[str, Any], override_logo: str | None) -> Path:
    logo_value = override_logo or str(spec["branding"]["logo_path"])
    candidate = Path(logo_value)
    possible_paths: list[Path] = []

    if candidate.is_absolute():
        possible_paths.append(candidate)
    else:
        possible_paths.append(Path.cwd() / candidate)
        possible_paths.append(spec_path.parent / candidate)
        possible_paths.append(spec_path.parent.parent / candidate)

    for path in possible_paths:
        if path.exists() and path.is_file():
            return path.resolve()

    checked_paths = ", ".join(str(path) for path in possible_paths)
    raise FileNotFoundError(
        "Logo-pad niet gevonden. "
        f"Ontvangen: '{logo_value}'. Gecontroleerde paden: {checked_paths}"
    )


def set_run_style(
    run: Any,
    *,
    font_name: str,
    size_pt: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_cell_background(cell: Any, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for old_node in tc_pr.findall(qn("w:shd")):
        tc_pr.remove(old_node)
    shade = OxmlElement("w:shd")
    shade.set(qn("w:val"), "clear")
    shade.set(qn("w:color"), "auto")
    shade.set(qn("w:fill"), clean_hex(hex_color, "#FFFFFF"))
    tc_pr.append(shade)


def configure_document_styles(doc: Document, spec: dict[str, Any]) -> None:
    branding = spec["branding"]
    layout = spec["document_layout"]

    margin_cm = float(layout.get("margin_cm", 2.0))
    section = doc.sections[0]
    section.top_margin = Cm(margin_cm)
    section.bottom_margin = Cm(margin_cm)
    section.left_margin = Cm(margin_cm)
    section.right_margin = Cm(margin_cm)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = branding["body_font"]
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), branding["body_font"])
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = hex_to_rgb(branding["text_color"], "#1f332c")


def add_document_header(doc: Document, spec: dict[str, Any], logo_path: Path) -> None:
    metadata = spec["metadata"]
    branding = spec["branding"]
    primary_color = hex_to_rgb(branding["primary_color"], "#17372c")
    secondary_color = hex_to_rgb(branding["secondary_color"], "#2d5647")
    accent_color = hex_to_rgb(branding["accent_color"], "#c99642")
    text_color = hex_to_rgb(branding["text_color"], "#1f332c")

    logo_paragraph = doc.add_paragraph()
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_run = logo_paragraph.add_run()
    logo_run.add_picture(str(logo_path), width=Cm(3.0))

    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_paragraph.add_run(metadata["title"])
    set_run_style(
        title_run,
        font_name=branding["heading_font"],
        size_pt=24,
        bold=True,
        color=primary_color,
    )

    subtitle_paragraph = doc.add_paragraph()
    subtitle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_paragraph.add_run(metadata["subtitle"])
    set_run_style(
        subtitle_run,
        font_name=branding["body_font"],
        size_pt=11,
        italic=True,
        color=secondary_color,
    )

    info_paragraph = doc.add_paragraph()
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info_paragraph.add_run(
        f"{metadata['practice_name']} | {metadata['practitioner']} | "
        f"{metadata['phone']} | {metadata['email']}"
    )
    set_run_style(
        info_run,
        font_name=branding["body_font"],
        size_pt=10,
        color=text_color,
    )

    note_paragraph = doc.add_paragraph()
    note_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note_paragraph.add_run(
        "Dit formulier is vertrouwelijk en bedoeld voor intake en veilige behandelaanpassing."
    )
    set_run_style(
        note_run,
        font_name=branding["body_font"],
        size_pt=9.5,
        color=accent_color,
        italic=True,
    )

    divider = doc.add_paragraph()
    divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_run = divider.add_run("".ljust(72, "_"))
    set_run_style(
        divider_run,
        font_name=branding["body_font"],
        size_pt=9,
        color=secondary_color,
    )

    doc.add_paragraph()


def add_section_title_row(table: Any, section_title: str, branding: dict[str, Any]) -> None:
    title_cell = table.rows[0].cells[0]
    set_cell_background(title_cell, branding["primary_color"])
    paragraph = title_cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(section_title)
    set_run_style(
        run,
        font_name=branding["heading_font"],
        size_pt=13,
        bold=True,
        color=RGBColor(255, 255, 255),
    )


def add_section_description_row(table: Any, description: str, branding: dict[str, Any]) -> None:
    if not description:
        return
    row = table.add_row()
    cell = row.cells[0]
    set_cell_background(cell, branding["surface_color"])
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(description)
    set_run_style(
        run,
        font_name=branding["body_font"],
        size_pt=10,
        italic=True,
        color=hex_to_rgb(branding["secondary_color"], "#2d5647"),
    )


def add_label_row(
    table: Any,
    *,
    label: str,
    required: bool,
    help_text: str,
    branding: dict[str, Any],
) -> None:
    row = table.add_row()
    cell = row.cells[0]
    set_cell_background(cell, branding["surface_color"])
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(1)

    main_run = paragraph.add_run(label)
    set_run_style(
        main_run,
        font_name=branding["body_font"],
        size_pt=10.5,
        bold=True,
        color=hex_to_rgb(branding["text_color"], "#1f332c"),
    )

    if required:
        required_run = paragraph.add_run(" *")
        set_run_style(
            required_run,
            font_name=branding["body_font"],
            size_pt=10.5,
            bold=True,
            color=hex_to_rgb(branding["accent_color"], "#c99642"),
        )

    if help_text:
        help_run = paragraph.add_run(f"\n{help_text}")
        set_run_style(
            help_run,
            font_name=branding["body_font"],
            size_pt=9,
            italic=True,
            color=hex_to_rgb(branding["secondary_color"], "#2d5647"),
        )


def add_marker_row(table: Any, marker_text: str, branding: dict[str, Any], *, left_indent_cm: float = 0.0) -> None:
    row = table.add_row()
    cell = row.cells[0]
    set_cell_background(cell, "#FFFFFF")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(1)
    paragraph.paragraph_format.space_after = Pt(4)
    if left_indent_cm:
        paragraph.paragraph_format.left_indent = Cm(left_indent_cm)
    marker_run = paragraph.add_run(marker_text)
    set_run_style(
        marker_run,
        font_name=branding["body_font"],
        size_pt=10.5,
        color=hex_to_rgb(branding["text_color"], "#1f332c"),
    )


def build_base_document(
    spec: dict[str, Any],
    logo_path: Path,
    output_path: Path,
) -> list[PlaceholderSpec]:
    doc = Document()
    configure_document_styles(doc, spec)
    add_document_header(doc, spec, logo_path)

    placeholders: list[PlaceholderSpec] = []
    branding = spec["branding"]
    section_spacing = float(spec["document_layout"].get("section_spacing_pt", 18))

    for section in spec["sections"]:
        table = doc.add_table(rows=1, cols=1)
        table.style = "Table Grid"
        table.autofit = True
        add_section_title_row(table, section["title"], branding)
        add_section_description_row(table, str(section.get("description", "")), branding)

        for field_spec in section["fields"]:
            field_type = field_spec["type"]
            field_id = field_spec["id"]
            label = field_spec["label"]
            required = bool(field_spec["required"])
            help_text = field_spec["help_text"]

            if field_type == "info_block":
                add_label_row(
                    table,
                    label=label,
                    required=False,
                    help_text=help_text,
                    branding=branding,
                )
                continue

            if field_type == "checkbox_group":
                add_label_row(
                    table,
                    label=label,
                    required=required,
                    help_text=help_text,
                    branding=branding,
                )
                for option in field_spec["options"]:
                    marker = marker_for(option["id"])
                    add_marker_row(
                        table,
                        f"{marker} {option['label']}",
                        branding,
                        left_indent_cm=0.35,
                    )
                    placeholders.append(
                        PlaceholderSpec(
                            marker=marker,
                            field_type="checkbox",
                            field_id=option["id"],
                            label=option["label"],
                            required=required,
                        )
                    )
                continue

            if field_type == "checkbox":
                marker = marker_for(field_id)
                add_marker_row(
                    table,
                    f"{marker} {label}{' *' if required else ''}",
                    branding,
                )
                if help_text:
                    add_label_row(
                        table,
                        label="Toelichting",
                        required=False,
                        help_text=help_text,
                        branding=branding,
                    )
                placeholders.append(
                    PlaceholderSpec(
                        marker=marker,
                        field_type=field_type,
                        field_id=field_id,
                        label=label,
                        required=required,
                    )
                )
                continue

            add_label_row(
                table,
                label=label,
                required=required,
                help_text=help_text,
                branding=branding,
            )

            marker = marker_for(field_id)
            add_marker_row(table, marker, branding)

            options = [str(option) for option in field_spec.get("options", [])]
            placeholders.append(
                PlaceholderSpec(
                    marker=marker,
                    field_type=field_type,
                    field_id=field_id,
                    label=label,
                    required=required,
                    options=options,
                )
            )

        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(section_spacing)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return placeholders


def replace_marker_with_control(word_doc: Any, placeholder: PlaceholderSpec) -> int:
    replaced_count = 0
    cursor = 0

    while cursor <= word_doc.Content.End:
        search_range = word_doc.Range(cursor, word_doc.Content.End)
        finder = search_range.Find
        finder.ClearFormatting()
        finder.Text = placeholder.marker
        finder.Forward = True
        finder.Wrap = WD_FIND_STOP

        if not finder.Execute():
            break

        found_range = word_doc.Range(search_range.Start, search_range.End)
        next_cursor = max(found_range.End, cursor + 1)
        replace_single_marker(word_doc, found_range, placeholder)
        replaced_count += 1
        cursor = next_cursor

    return replaced_count


def replace_single_marker(word_doc: Any, found_range: Any, placeholder: PlaceholderSpec) -> None:
    field_type = placeholder.field_type

    if field_type == "signature_line":
        found_range.Text = "______________________________ (Handtekening)"
        return

    start = found_range.Start
    found_range.Text = ""
    insertion_range = word_doc.Range(start, start)

    if field_type == "text_short":
        control = word_doc.ContentControls.Add(WD_CONTENT_CONTROL_TEXT, insertion_range)
        control.Title = placeholder.label
        control.Tag = placeholder.field_id
        try:
            control.MultiLine = False
        except Exception:
            pass
        return

    if field_type == "text_long":
        control = word_doc.ContentControls.Add(WD_CONTENT_CONTROL_RICH_TEXT, insertion_range)
        control.Title = placeholder.label
        control.Tag = placeholder.field_id
        try:
            control.MultiLine = True
        except Exception:
            pass
        return

    if field_type == "checkbox":
        control = word_doc.ContentControls.Add(WD_CONTENT_CONTROL_CHECKBOX, insertion_range)
        control.Title = placeholder.label
        control.Tag = placeholder.field_id
        try:
            control.Checked = False
        except Exception:
            pass
        return

    if field_type == "dropdown":
        control = word_doc.ContentControls.Add(WD_CONTENT_CONTROL_DROPDOWN, insertion_range)
        control.Title = placeholder.label
        control.Tag = placeholder.field_id
        entries = control.DropdownListEntries
        try:
            while entries.Count > 0:
                entries.Item(1).Delete()
        except Exception:
            pass
        entries.Add("Kies een optie", f"{placeholder.field_id}_0", 1)
        for index, option in enumerate(placeholder.options, start=2):
            entries.Add(str(option), f"{placeholder.field_id}_{index}", index)
        return

    if field_type == "date":
        control = word_doc.ContentControls.Add(WD_CONTENT_CONTROL_DATE, insertion_range)
        control.Title = placeholder.label
        control.Tag = placeholder.field_id
        try:
            control.DateDisplayFormat = "dd-MM-yyyy"
            control.DateDisplayLocale = 1043
        except Exception:
            pass
        return

    raise RuntimeError(f"Onbekend field_type voor COM-replacement: {field_type}")


def apply_com_controls(
    base_doc_path: Path,
    output_path: Path,
    placeholders: list[PlaceholderSpec],
) -> None:
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        raise RuntimeError(
            "pywin32/Word COM is niet beschikbaar. "
            "Installeer pywin32 of gebruik --no-com voor fallback-output."
        ) from exc

    word_app = None
    word_doc = None
    replacement_errors: list[str] = []

    try:
        word_app = win32com.client.DispatchEx("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0
        word_doc = word_app.Documents.Open(str(base_doc_path.resolve()))

        for placeholder in placeholders:
            count = replace_marker_with_control(word_doc, placeholder)
            if count == 0:
                replacement_errors.append(
                    f"Marker niet gevonden in document: {placeholder.marker}"
                )
            elif count > 1:
                replacement_errors.append(
                    f"Marker kwam meerdere keren voor ({count}x): {placeholder.marker}"
                )

        if replacement_errors:
            raise RuntimeError(
                "Niet alle markers konden correct worden vervangen:\n"
                + "\n".join(replacement_errors)
            )

        output_path.parent.mkdir(parents=True, exist_ok=True)
        word_doc.SaveAs2(str(output_path.resolve()))

    finally:
        if word_doc is not None:
            try:
                word_doc.Close(False)
            except Exception:
                pass
        if word_app is not None:
            try:
                word_app.Quit()
            except Exception:
                pass


def iter_all_paragraphs(doc: Document) -> Iterable[Any]:
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def apply_plain_fallback(base_doc_path: Path, placeholders: list[PlaceholderSpec]) -> None:
    replacement_map: dict[str, str] = {}
    for placeholder in placeholders:
        if placeholder.field_type == "checkbox":
            replacement_map[placeholder.marker] = "[ ]"
        elif placeholder.field_type == "date":
            replacement_map[placeholder.marker] = "__-__-____"
        elif placeholder.field_type == "dropdown":
            replacement_map[placeholder.marker] = "[Kies een optie]"
        elif placeholder.field_type == "signature_line":
            replacement_map[placeholder.marker] = "______________________________ (Handtekening)"
        elif placeholder.field_type == "text_long":
            replacement_map[placeholder.marker] = "____________________________________________"
        else:
            replacement_map[placeholder.marker] = "____________________________"

    doc = Document(str(base_doc_path))
    for paragraph in iter_all_paragraphs(doc):
        updated_text = paragraph.text
        for marker, replacement in replacement_map.items():
            if marker in updated_text:
                updated_text = updated_text.replace(marker, replacement)
        if updated_text != paragraph.text:
            paragraph.text = updated_text
    doc.save(str(base_doc_path))


def run_generation(
    spec_path: Path,
    output_path: Path,
    override_logo: str | None,
    validate_only: bool,
    no_com: bool,
) -> None:
    spec = load_spec(spec_path)
    validate_spec(spec)
    logo_path = resolve_logo_path(spec_path, spec, override_logo)

    if validate_only:
        print(f"Validatie geslaagd voor spec: {spec_path}")
        print(f"Gevonden logo: {logo_path}")
        return

    with tempfile.TemporaryDirectory(prefix="menyentuh_form_") as tmp_dir:
        base_doc_path = Path(tmp_dir) / "base_form.docx"
        placeholders = build_base_document(spec, logo_path, base_doc_path)

        if no_com:
            apply_plain_fallback(base_doc_path, placeholders)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copyfile(base_doc_path, output_path)
            print(
                "Waarschuwing: '--no-com' actief. Output bevat geen echte interactieve Word-controls."
            )
            print(f"Document opgeslagen op: {output_path}")
            return

        apply_com_controls(base_doc_path, output_path, placeholders)
        print(f"Document met interactieve velden opgeslagen op: {output_path}")


def main() -> int:
    args = parse_args()
    spec_path = Path(args.spec).resolve()
    output_path = Path(args.out).resolve()

    try:
        run_generation(
            spec_path=spec_path,
            output_path=output_path,
            override_logo=args.logo,
            validate_only=args.validate_only,
            no_com=args.no_com,
        )
        return 0
    except (SpecValidationError, FileNotFoundError, RuntimeError, OSError) as exc:
        print(f"Fout: {exc}", file=sys.stderr)
        return 1
    except Exception as exc:
        print(f"Onverwachte fout: {exc}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    raise SystemExit(main())
